import json
import requests
import time
from typing import List, Dict, Any, Optional, Union, Tuple, Callable, Iterator
import uuid
import asyncio
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime

class Message:
    """A class representing a chat message."""
    
    def __init__(self, role: str, content: str):
        """
        Initialize a chat message.
        
        Args:
            role: The role of the sender (e.g., 'system', 'user', 'assistant')
            content: The content of the message
        """
        self.role = role
        self.content = content
        
    def to_dict(self) -> Dict[str, str]:
        """Convert the message to a dictionary representation."""
        return {"role": self.role, "content": self.content}
    
    @classmethod
    def from_dict(cls, message_dict: Dict[str, str]) -> 'Message':
        """Create a Message from a dictionary representation."""
        return cls(role=message_dict["role"], content=message_dict["content"])
    
    def __str__(self) -> str:
        return f"{self.role}: {self.content}"


class ChatOllama:
    """Chat model wrapper for Ollama API"""
    
    def __init__(
        self,
        model: str = "llama2",
        base_url: str = "http://localhost:11434",
        temperature: float = 0.7,
        top_p: float = 0.9,
        top_k: int = 40,
        max_tokens: int = 1024,
        system: Optional[str] = None,
        timeout: Optional[float] = 120,
        keep_alive: Optional[str] = None,
        format: Optional[str] = None,
        context: Optional[List[int]] = None,
        options: Optional[Dict[str, Any]] = None,
    ):
        """
        Initialize ChatOllama model.
        
        Args:
            model: Name of the model to use
            base_url: Base URL of the Ollama API
            temperature: Sampling temperature (0.0-1.0)
            top_p: Top-p sampling (nucleus sampling)
            top_k: Top-k sampling
            max_tokens: Maximum number of tokens to generate
            system: Optional system prompt to use with all requests
            timeout: Request timeout in seconds
            keep_alive: Optional keep alive time for model in seconds or minutes
            format: Optional format for responses (json, etc.)
            context: Optional context window for model
            options: Additional options to pass to the Ollama API
        """
        self.model = model
        self.base_url = base_url.rstrip('/')
        self.temperature = temperature
        self.top_p = top_p
        self.top_k = top_k
        self.max_tokens = max_tokens
        self.system = system
        self.timeout = timeout
        self.keep_alive = keep_alive
        self.format = format
        self.context = context
        self.options = options or {}
        self._validate_params()
        
    def _validate_params(self):
        """Validate parameters are within valid ranges"""
        if self.temperature < 0.0 or self.temperature > 1.0:
            raise ValueError("Temperature must be between 0.0 and 1.0")
        if self.top_p < 0.0 or self.top_p > 1.0:
            raise ValueError("Top-p must be between 0.0 and 1.0")
        if self.top_k < 1:
            raise ValueError("Top-k must be at least 1")
        if self.max_tokens < 1:
            raise ValueError("Max tokens must be at least 1")
    
    def _prepare_messages(self, messages: List[Union[Dict[str, str], Message]]) -> List[Dict[str, str]]:
        """Convert messages to the format expected by Ollama API"""
        normalized_messages = []
        
        for message in messages:
            if isinstance(message, dict):
                normalized_messages.append(message)
            elif isinstance(message, Message):
                normalized_messages.append(message.to_dict())
            else:
                raise ValueError(f"Unsupported message type: {type(message)}")
                
        # Add system message if provided and not already in messages
        if self.system and not any(msg.get("role") == "system" for msg in normalized_messages):
            normalized_messages.insert(0, {"role": "system", "content": self.system})
            
        return normalized_messages
    
    def _prepare_request(self, messages: List[Dict[str, str]], stream: bool = False) -> Dict[str, Any]:
        """Prepare the request body for the Ollama API"""
        # Basic parameters
        request = {
            "model": self.model,
            "messages": messages,
            "stream": stream,
            "options": {
                "temperature": self.temperature,
                "top_p": self.top_p,
                "top_k": self.top_k,
                "num_predict": self.max_tokens,
            }
        }
        
        # Add format if specified
        if self.format:
            request["format"] = self.format
            
        # Add context if specified
        if self.context:
            request["context"] = self.context
            
        # Add keep_alive if specified
        if self.keep_alive:
            request["keep_alive"] = self.keep_alive
            
        # Add any additional options
        if self.options:
            request["options"].update(self.options)
            
        return request
    
    def _make_chat_request(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        """Make a request to the Ollama chat API endpoint"""
        try:
            response = requests.post(
                f"{self.base_url}/api/chat",
                json=payload,
                timeout=self.timeout,
            )
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            error_msg = str(e)
            try:
                if hasattr(e, "response") and e.response is not None:
                    error_details = e.response.json()
                    error_msg = f"{error_msg}: {error_details.get('error', '')}"
            except:
                pass
            raise RuntimeError(f"Ollama API request failed: {error_msg}")
    
    def _stream_chat_request(self, payload: Dict[str, Any]) -> Iterator[Dict[str, Any]]:
        """Stream responses from the Ollama chat API endpoint"""
        try:
            with requests.post(
                f"{self.base_url}/api/chat",
                json=payload,
                timeout=self.timeout,
                stream=True,
            ) as response:
                response.raise_for_status()
                for line in response.iter_lines():
                    if line:
                        yield json.loads(line)
        except requests.exceptions.RequestException as e:
            error_msg = str(e)
            try:
                if hasattr(e, "response") and e.response is not None:
                    error_details = e.response.json()
                    error_msg = f"{error_msg}: {error_details.get('error', '')}"
            except:
                pass
            raise RuntimeError(f"Ollama API stream request failed: {error_msg}")
    
    def chat(self, messages: List[Union[Dict[str, str], Message]]) -> Dict[str, Any]:
        """
        Get a chat response from the model.
        
        Args:
            messages: List of message dictionaries or Message objects
            
        Returns:
            Response from the model with assistant message
        """
        normalized_messages = self._prepare_messages(messages)
        payload = self._prepare_request(normalized_messages, stream=False)
        return self._make_chat_request(payload)
    
    def stream_chat(self, messages: List[Union[Dict[str, str], Message]]) -> Iterator[Dict[str, Any]]:
        """
        Stream a chat response from the model.
        
        Args:
            messages: List of message dictionaries or Message objects
            
        Returns:
            Iterator of streaming responses
        """
        normalized_messages = self._prepare_messages(messages)
        payload = self._prepare_request(normalized_messages, stream=True)
        return self._stream_chat_request(payload)
    
    async def achat(self, messages: List[Union[Dict[str, str], Message]]) -> Dict[str, Any]:
        """
        Async version of chat.
        
        Args:
            messages: List of message dictionaries or Message objects
            
        Returns:
            Response from the model with assistant message
        """
        normalized_messages = self._prepare_messages(messages)
        payload = self._prepare_request(normalized_messages, stream=False)
        
        with ThreadPoolExecutor() as executor:
            return await asyncio.get_event_loop().run_in_executor(
                executor, self._make_chat_request, payload
            )
    
    def generate(self, prompt: str) -> str:
        """
        Get a response for a single prompt.
        
        Args:
            prompt: String prompt
            
        Returns:
            Model response as a string
        """
        messages = [{"role": "user", "content": prompt}]
        response = self.chat(messages)
        return response.get("message", {}).get("content", "")
    
    def __call__(self, messages: List[Union[Dict[str, str], Message]]) -> str:
        """
        Call the model directly with a list of messages.
        
        Args:
            messages: List of message dictionaries or Message objects
            
        Returns:
            Model response as a string
        """
        response = self.chat(messages)
        return response.get("message", {}).get("content", "")


class OllamaEmbeddings:
    """Embeddings implementation using Ollama API"""
    
    def __init__(
        self,
        model: str = "llama2",
        base_url: str = "http://localhost:11434",
        timeout: Optional[float] = 120,
        options: Optional[Dict[str, Any]] = None,
        dimensions: Optional[int] = None,
    ):
        """
        Initialize Ollama embeddings.
        
        Args:
            model: Name of the model to use
            base_url: Base URL of the Ollama API
            timeout: Request timeout in seconds
            options: Additional options to pass to the Ollama API
            dimensions: Optional dimension for embeddings
        """
        self.model = model
        self.base_url = base_url.rstrip('/')
        self.timeout = timeout
        self.options = options or {}
        self.dimensions = dimensions
        
        if dimensions is not None:
            self.options["dimensions"] = dimensions
    
    def _make_embed_request(self, text: str) -> Dict[str, Any]:
        """Make a request to the Ollama embeddings API endpoint"""
        payload = {
            "model": self.model,
            "prompt": text,
        }
        
        if self.options:
            payload["options"] = self.options
        
        try:
            response = requests.post(
                f"{self.base_url}/api/embeddings",
                json=payload,
                timeout=self.timeout,
            )
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            error_msg = str(e)
            try:
                if hasattr(e, "response") and e.response is not None:
                    error_details = e.response.json()
                    error_msg = f"{error_msg}: {error_details.get('error', '')}"
            except:
                pass
            raise RuntimeError(f"Ollama API embeddings request failed: {error_msg}")
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """
        Get embeddings for a list of documents.
        
        Args:
            texts: List of document texts
            
        Returns:
            List of embeddings for each document
        """
        embeddings = []
        
        for text in texts:
            response = self._make_embed_request(text)
            embeddings.append(response.get("embedding", []))
            
        return embeddings
    
    def embed_query(self, text: str) -> List[float]:
        """
        Get embedding for a query text.
        
        Args:
            text: Query text
            
        Returns:
            Embedding for the query
        """
        response = self._make_embed_request(text)
        return response.get("embedding", [])
    
    async def aembed_documents(self, texts: List[str]) -> List[List[float]]:
        """
        Async version of embed_documents.
        
        Args:
            texts: List of document texts
            
        Returns:
            List of embeddings for each document
        """
        async def _embed_single(text):
            with ThreadPoolExecutor() as executor:
                response = await asyncio.get_event_loop().run_in_executor(
                    executor, self._make_embed_request, text
                )
                return response.get("embedding", [])
        
        return await asyncio.gather(*[_embed_single(text) for text in texts])
    
    async def aembed_query(self, text: str) -> List[float]:
        """
        Async version of embed_query.
        
        Args:
            text: Query text
            
        Returns:
            Embedding for the query
        """
        with ThreadPoolExecutor() as executor:
            response = await asyncio.get_event_loop().run_in_executor(
                executor, self._make_embed_request, text
            )
            return response.get("embedding", [])
    
    def get_dimension(self) -> int:
        """
        Get the dimension of the embeddings.
        
        Returns:
            Dimension of the embeddings
        """
        # If dimensions were explicitly set, return that value
        if self.dimensions is not None:
            return self.dimensions
            
        # Otherwise, get a sample embedding to determine the dimension
        try:
            sample = self.embed_query("Sample text")
            return len(sample)
        except:
            # If we can't get a sample, return a default dimension
            return 4096  # Common dimension for many Ollama models


# Usage example:
# chat_model = ChatOllama(model="smollm")
# response = chat_model.chat([
#     {"role": "user", "content": "Hello, how are you?"}
# ])
# print(response["message"]["content"])

embeddings_model = OllamaEmbeddings(model="mxbai-embed-large")
# doc_embeddings = embeddings_model.aembed_documents(["This is a document"])
# query_embedding = embeddings_model.aembed_query("This is a query")

# print(doc_embeddings)
# print(query_embedding)
# <coroutine object OllamaEmbeddings.aembed_documents at 0x00000247B35D4F20>
# <coroutine object OllamaEmbeddings.aembed_query at 0x00000247B35C6F40>
#display the content of the embeddings
# Option 1: Using async/await in an async function
# Option 1: Using async/await in an async function
async def main():
    start_time = time.time()
    doc_embeddings = await embeddings_model.aembed_documents(["This is a document"])
    query_embedding = await embeddings_model.aembed_query("This is a query")
    end_time = time.time()
    
    print(f"Option 1 time taken: {end_time - start_time:.2f} seconds")
    # print("Doc embeddings:", doc_embeddings)
    # print("Query embedding:", query_embedding)


    # print("Doc embeddings:", doc_embeddings)
    # print("Query embedding:", query_embedding)

    # Option 3: Using asyncio.gather()
    start_time = time.time()
    doc_embeddings, query_embedding = await asyncio.gather(
        embeddings_model.aembed_documents(["This is a document"]),
        embeddings_model.aembed_query("This is a query"),
    )
    end_time = time.time()
    
    print(f"\nOption 3 time taken: {end_time - start_time:.2f} seconds")
    # print("Doc embeddings:", doc_embeddings)
    # print("Query embedding:", query_embedding)


        # print("Doc embeddings:", doc_embeddings)
        # print("Query embedding:", query_embedding)

# asyncio.run(main())
# Option 1 time taken: 4.24 seconds
# Option 3 time taken: 2.17 seconds




txt="""
This is a sample text file to test LangChain function alternatives.
LangChain is a framework for developing applications powered by language models.

Here are some example sentences:
1. The quick brown fox jumps over the lazy dog.
2. Artificial intelligence is transforming the world.
3. Natural language processing enables machines to understand human language.
4. Machine learning algorithms can identify patterns in data.
5. Data science combines domain expertise, programming skills, and knowledge of mathematics and statistics.
6. Cloud computing provides scalable resources over the internet.
7. Cybersecurity is essential to protect data and systems from threats.
8. Blockchain technology ensures secure and transparent transactions.
9. The Internet of Things connects devices and enables smart environments.
10. Augmented reality overlays digital information onto the real world.
11. Virtual reality creates immersive digital experiences.
12. Quantum computing leverages quantum mechanics to solve complex problems.
13. Robotics integrates engineering and computer science to build automated machines.
14. Autonomous vehicles use sensors and AI to navigate without human intervention.
15. Big data analytics extracts insights from large datasets.
16. Edge computing processes data closer to where it is generated.
17. 5G technology offers faster and more reliable wireless communication.
18. Bioinformatics applies computational techniques to biological data.
19. Fintech innovations are transforming the financial industry.
20. E-commerce platforms facilitate online buying and selling.
21. Social media connects people and enables information sharing.
22. Digital marketing leverages online channels to reach customers.
23. User experience design focuses on creating intuitive interfaces.
24. Software development methodologies include Agile and DevOps.
25. Version control systems like Git manage code changes.
26. Continuous integration and continuous deployment automate software delivery.
27. Containerization with Docker ensures consistent environments.
28. Microservices architecture breaks applications into smaller, independent services.
29. RESTful APIs enable communication between different systems.
30. GraphQL provides a flexible approach to querying APIs.
31. Serverless computing abstracts infrastructure management.
32. DevSecOps integrates security practices into the development process.
33. Test-driven development emphasizes writing tests before code.
34. Behavior-driven development focuses on user behavior and requirements.
35. Pair programming involves two developers working together on the same code.
36. Code reviews ensure quality and share knowledge among team members.
37. Refactoring improves code structure without changing its behavior.
38. Technical debt represents the cost of maintaining suboptimal code.
39. Design patterns provide reusable solutions to common problems.
40. Object-oriented programming organizes code into classes and objects.
41. Functional programming treats computation as the evaluation of mathematical functions.
42. Procedural programming follows a sequence of instructions.
43. Scripting languages automate repetitive tasks.
44. Markup languages like HTML structure web content.
45. Cascading Style Sheets (CSS) control the presentation of web pages.
46. JavaScript adds interactivity to web pages.
47. Front-end frameworks like React and Angular build dynamic user interfaces.
48. Back-end frameworks like Django and Flask handle server-side logic.
49. Databases store and manage data for applications.
50. SQL queries retrieve and manipulate data in relational databases.
51. NoSQL databases handle unstructured data.
52. Data warehousing consolidates data from multiple sources.
53. Data lakes store raw data in its native format.
54. ETL processes extract, transform, and load data.
55. Data visualization tools like Tableau and Power BI present data insights.
56. Business intelligence systems support decision-making.
57. Machine learning models are trained on data to make predictions.
58. Deep learning uses neural networks with many layers.
59. Natural language generation creates human-like text.
60. Speech recognition converts spoken language into text.
61. Computer vision interprets visual information from images and videos.
62. Reinforcement learning trains agents through trial and error.
63. Transfer learning applies knowledge from one domain to another.
64. Hyperparameter tuning optimizes machine learning models.
65. Model deployment makes machine learning models available for use.
66. A/B testing compares two versions of a system to determine which performs better.
67. User feedback informs product improvements.
68. Accessibility ensures that applications are usable by everyone.
69. Localization adapts software for different languages and regions.
70. Internationalization designs software to support multiple languages and regions.
71. Cloud-native applications are built to run in cloud environments.
72. Infrastructure as Code manages infrastructure through code.
73. Configuration management tools like Ansible and Puppet automate system setup.
74. Monitoring and logging track system performance and issues.
75. Incident response addresses and resolves system failures.
76. Disaster recovery plans prepare for data loss and system outages.
77. Load balancing distributes traffic across multiple servers.
78. Content delivery networks (CDNs) accelerate the delivery of web content.
79. API gateways manage and secure API traffic.
80. Identity and access management controls user permissions.
81. Single sign-on (SSO) allows users to access multiple systems with one login.
82. Multi-factor authentication (MFA) adds an extra layer of security.
83. Encryption protects data in transit and at rest.
84. Data privacy regulations like GDPR and CCPA protect user information.
85. Ethical AI considers the societal impact of artificial intelligence.
86. Open source software is developed collaboratively and shared freely.
87. Licensing agreements define how software can be used and distributed.
88. Intellectual property rights protect software innovations.
89. Software patents grant exclusive rights to inventors.
90. Copyright law protects original works of authorship.
91. Trademarks protect brand names and logos.
92. Compliance ensures adherence to laws and regulations.
93. Risk management identifies and mitigates potential issues.
94. Project management methodologies include Waterfall and Scrum.
95. Stakeholder engagement involves communicating with those affected by a project.
96. Change management addresses the human side of organizational change.
97. Knowledge management captures and shares organizational knowledge.
98. Professional development supports continuous learning and growth.
99. Mentorship programs pair experienced professionals with less experienced colleagues.
100. Networking builds relationships within and outside the organization.

End of sample text."""

import sys


from dataclasses import dataclass
from typing import Dict, Any
import json
# import psycopg2
import numpy as np
# from psycopg2.extras import execute_values
# import faiss
import chromadb
import asyncio
from concurrent.futures import ThreadPoolExecutor
import os
@dataclass
class Document:
    page_content: str
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
    

class VectorStore:
    """Stores and retrieves documents using vector embeddings"""
    
    def __init__(self, 
                 embedding_model=None,
                 store_type: str = "in_memory",
                 connection_string: str = None,
                 collection_name: str = "documents",
                 dimension: int = 1536):  # Default for OpenAI embeddings
        """
        Initialize a vector store
        
        Args:
            embedding_model: Model to create embeddings
            store_type: Type of vector store ('in_memory', 'faiss', 'chroma', etc.)
            connection_string: Connection string for external stores
            collection_name: Name of collection for storage
            dimension: Embedding dimension
        """
        self.embedding_model = embedding_model
        self.store_type = store_type
        self.connection_string = connection_string
        self.collection_name = collection_name
        self.dimension = dimension
        
        # Storage for in-memory vector store
        self.documents = []
        self.embeddings = []
        self.metadatas = []
        self.ids = []
        
        # Initialize the appropriate vector store
        self._initialize_store()
    
    def _initialize_store(self):
        """Initialize the vector store based on store_type"""
        if self.store_type == "in_memory":
            # No additional initialization needed
            pass
            
        elif self.store_type == "faiss":
            try:
                import faiss
                import numpy as np
                
                # Initialize FAISS index
                self.index = faiss.IndexFlatL2(self.dimension)
                self.vector_id = 0  # Counter for document IDs
                
            except ImportError:
                raise ImportError("FAISS store requires 'faiss-cpu' or 'faiss-gpu' package")
                
        elif self.store_type == "chroma":
            try:
                import chromadb
                
                # Initialize ChromaDB client
                if self.connection_string:
                    self.client = chromadb.Client(self.connection_string)
                else:
                    self.client = chromadb.Client()
                    
                # Get or create collection
                self.collection = self.client.get_or_create_collection(
                    name=self.collection_name,
                    embedding_function=self._get_embeddings if self.embedding_model else None
                )
                
            except ImportError:
                raise ImportError("Chroma store requires 'chromadb' package")
                
        elif self.store_type == "postgres":
            try:
                import psycopg2
                import numpy as np
                from psycopg2.extras import execute_values
                
                # Connect to PostgreSQL
                self.conn = psycopg2.connect(self.connection_string)
                self.cursor = self.conn.cursor()
                
                # Create tables if they don't exist
                self.cursor.execute("""
                CREATE TABLE IF NOT EXISTS documents (
                    id SERIAL PRIMARY KEY,
                    content TEXT NOT NULL,
                    metadata JSONB,
                    embedding_id INTEGER
                );
                
                CREATE TABLE IF NOT EXISTS embeddings (
                    id SERIAL PRIMARY KEY,
                    vector REAL[] NOT NULL
                );
                """)
                
                self.conn.commit()
                
                # Create vector extension if not exists
                try:
                    self.cursor.execute("CREATE EXTENSION IF NOT EXISTS vector;")
                    self.conn.commit()
                except:
                    print("Warning: vector extension not available in PostgreSQL")
                
            except ImportError:
                raise ImportError("PostgreSQL store requires 'psycopg2' package")
                
        else:
            raise ValueError(f"Unsupported store type: {self.store_type}")
    
    def _get_embeddings(self, texts):
        """Get embeddings for texts using the embedding model"""
        if not self.embedding_model:
            raise ValueError("No embedding model provided")
            
        embeddings = []
        
        # Handle different embedding model interfaces
        for text in texts:
            if hasattr(self.embedding_model, "embed_query"):
                # LangChain interface
                embedding = self.embedding_model.embed_query(text)
            elif hasattr(self.embedding_model, "embed"):
                # Generic interface
                embedding = self.embedding_model.embed(text)
            elif hasattr(self.embedding_model, "__call__"):
                # Callable interface
                embedding = self.embedding_model(text)
            elif hasattr(self.embedding_model, "encode"):
                # Sentence Transformers interface
                embedding = self.embedding_model.encode(text)
            else:
                raise ValueError("Unsupported embedding model interface")
            
            embeddings.append(embedding)
            
        return embeddings
    
    def add_documents(self, documents):
        """
        Add documents to the vector store
        
        Args:
            documents: List of Document objects
            
        Returns:
            List of document IDs
        """
        if not documents:
            return []
            
        # Extract texts for embedding
        texts = [doc.page_content for doc in documents]
        
        # Extract metadatas
        metadatas = [doc.metadata for doc in documents]
        
        # Generate embeddings
        embeddings = self._get_embeddings(texts)
        
        # Add to store based on type
        if self.store_type == "in_memory":
            # Generate IDs
            ids = [str(len(self.documents) + i) for i in range(len(documents))]
            
            # Store in memory
            self.documents.extend(documents)
            self.embeddings.extend(embeddings)
            self.metadatas.extend(metadatas)
            self.ids.extend(ids)
            
            return ids
            
        elif self.store_type == "faiss":
            import numpy as np
            
            # Convert embeddings to numpy array
            embeddings_array = np.array(embeddings).astype('float32')
            
            # Generate IDs
            ids = [str(self.vector_id + i) for i in range(len(documents))]
            
            # Add vectors to FAISS index
            self.index.add(embeddings_array)
            
            # Store document data
            for i, doc in enumerate(documents):
                doc.metadata['faiss_id'] = ids[i]
                self.documents.append(doc)
                
            # Update ID counter
            self.vector_id += len(documents)
            
            return ids
            
        elif self.store_type == "chroma":
            # Generate IDs
            ids = [str(hash(doc.page_content))[:16] for doc in documents]
            
            # Add to ChromaDB
            self.collection.add(
                documents=texts,
                embeddings=embeddings,
                metadatas=metadatas,
                ids=ids
            )
            
            return ids
            
        elif self.store_type == "postgres":
            ids = []
            
            # Insert embeddings and documents
            for i, doc in enumerate(documents):
                # Insert embedding
                self.cursor.execute(
                    "INSERT INTO embeddings (vector) VALUES (%s) RETURNING id",
                    (embeddings[i],)
                )
                embedding_id = self.cursor.fetchone()[0]
                
                # Insert document
                self.cursor.execute(
                    "INSERT INTO documents (content, metadata, embedding_id) VALUES (%s, %s, %s) RETURNING id",
                    (doc.page_content, json.dumps(doc.metadata), embedding_id)
                )
                doc_id = self.cursor.fetchone()[0]
                
                ids.append(str(doc_id))
                
            self.conn.commit()
            
            return ids
    
    def similarity_search(self, query, k=4):
        """
        Find documents similar to the query
        
        Args:
            query: Query string
            k: Number of documents to retrieve
            
        Returns:
            List of retrieved Document objects
        """
        # Get query embedding
        query_embedding = self._get_embeddings([query])[0]
        
        # Search based on store type
        if self.store_type == "in_memory":
            import numpy as np
            
            if not self.embeddings:
                return []
                
            # Convert to numpy array for efficient computation
            embeddings_array = np.array(self.embeddings)
            query_array = np.array(query_embedding)
            
            # Compute similarities
            similarities = np.dot(embeddings_array, query_array) / (
                np.linalg.norm(embeddings_array, axis=1) * np.linalg.norm(query_array)
            )
            
            # Get top k indices
            top_k_indices = np.argsort(similarities)[-k:][::-1]
            
            # Get documents
            results = [self.documents[i] for i in top_k_indices]
            
            return results
            
        elif self.store_type == "faiss":
            import numpy as np
            
            if len(self.documents) == 0:
                return []
                
            # Convert query embedding to numpy array
            query_array = np.array([query_embedding]).astype('float32')
            
            # Search FAISS index
            distances, indices = self.index.search(query_array, k)
            
            # Get documents
            results = []
            for idx in indices[0]:
                if 0 <= idx < len(self.documents):
                    results.append(self.documents[idx])
                    
            return results
            
        elif self.store_type == "chroma":
            # Query ChromaDB
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=k
            )
            
            # Convert results to Document objects
            documents = []
            for i in range(len(results['ids'][0])):
                doc = Document(
                    page_content=results['documents'][0][i],
                    metadata=results['metadatas'][0][i] if results['metadatas'][0] else {}
                )
                documents.append(doc)
                
            return documents
            
        elif self.store_type == "postgres":
            # Get query vector
            query_vector = query_embedding
            
            try:
                # Try using vector similarity if extension is available
                self.cursor.execute("""
                    SELECT d.id, d.content, d.metadata FROM documents d
                    JOIN embeddings e ON d.embedding_id = e.id
                    ORDER BY e.vector <-> %s
                    LIMIT %s
                """, (query_vector, k))
            except:
                # Fallback to full scan if vector extension not available
                self.cursor.execute("""
                    SELECT d.id, d.content, d.metadata FROM documents d
                    JOIN embeddings e ON d.embedding_id = e.id
                    LIMIT %s
                """, (k,))
                
            results = self.cursor.fetchall()
            
            # Convert to Document objects
            documents = []
            for row in results:
                doc = Document(
                    page_content=row[1],
                    metadata=json.loads(row[2]) if row[2] else {}
                )
                documents.append(doc)
                
            return documents
    
    def delete(self, document_ids):
        """
        Delete documents from the vector store
        
        Args:
            document_ids: List of document IDs to delete
            
        Returns:
            Boolean indicating success
        """
        if not document_ids:
            return True
            
        # Handle deletion based on store type
        if self.store_type == "in_memory":
            # Find indices to delete
            indices_to_delete = []
            for i, doc_id in enumerate(self.ids):
                if doc_id in document_ids:
                    indices_to_delete.append(i)
                    
            # Delete in reverse order to avoid index shifting
            for i in sorted(indices_to_delete, reverse=True):
                del self.documents[i]
                del self.embeddings[i]
                del self.metadatas[i]
                del self.ids[i]
                
            return True
            
        elif self.store_type == "chroma":
            # Delete from ChromaDB
            self.collection.delete(ids=document_ids)
            return True
            
        elif self.store_type == "postgres":
            # Convert string IDs to integers
            int_ids = [int(doc_id) for doc_id in document_ids]
            
            # Delete from PostgreSQL
            self.cursor.execute("""
                DELETE FROM documents WHERE id = ANY(%s)
            """, (int_ids,))
            
            self.conn.commit()
            return True
            
        elif self.store_type == "faiss":
            # FAISS doesn't support direct deletion
            # We would need to rebuild the index, which we'll skip for simplicity
            print("Warning: Deletion not supported for FAISS vector store")
            return False
    
    def clear(self):
        """
        Clear all documents from the vector store
        
        Returns:
            Boolean indicating success
        """
        # Handle clearing based on store type
        if self.store_type == "in_memory":
            self.documents = []
            self.embeddings = []
            self.metadatas = []
            self.ids = []
            return True
            
        elif self.store_type == "faiss":
            import faiss
            
            # Reset FAISS index
            self.index = faiss.IndexFlatL2(self.dimension)
            self.documents = []
            self.vector_id = 0
            return True
            
        elif self.store_type == "chroma":
            # Clear ChromaDB collection
            self.collection.delete(ids=self.collection.get()["ids"])
            return True
            
        elif self.store_type == "postgres":
            # Truncate tables
            self.cursor.execute("TRUNCATE documents, embeddings")
            self.conn.commit()
            return True
    
    def get_collection_stats(self):
        """
        Get statistics about the vector store collection
        
        Returns:
            Dictionary with collection statistics
        """
        stats = {
            "store_type": self.store_type,
            "dimension": self.dimension,
            "collection_name": self.collection_name
        }
        
        if self.store_type == "in_memory":
            stats.update({
                "document_count": len(self.documents),
                "total_embeddings": len(self.embeddings),
                "memory_usage_estimate": sum(sys.getsizeof(emb) for emb in self.embeddings) + 
                                        sum(sys.getsizeof(doc.page_content) for doc in self.documents)
            })
            
        elif self.store_type == "faiss":
            stats.update({
                "document_count": len(self.documents),
                "index_type": type(self.index).__name__,
                "is_trained": getattr(self.index, "is_trained", True)
            })
            
            # Add FAISS-specific stats if available
            if hasattr(self.index, "ntotal"):
                stats["index_size"] = self.index.ntotal
                
        elif self.store_type == "chroma":
            # Get collection info from ChromaDB
            try:
                collection_data = self.collection.get()
                stats.update({
                    "document_count": len(collection_data.get("ids", [])),
                    "embedding_function": str(self.collection._embedding_function.__class__.__name__ 
                                            if hasattr(self.collection, "_embedding_function") else "None")
                })
            except Exception as e:
                stats["error"] = f"Error retrieving ChromaDB stats: {str(e)}"
                
        elif self.store_type == "postgres":
            try:
                # Count documents
                self.cursor.execute("SELECT COUNT(*) FROM documents")
                doc_count = self.cursor.fetchone()[0]
                
                # Count embeddings
                self.cursor.execute("SELECT COUNT(*) FROM embeddings")
                emb_count = self.cursor.fetchone()[0]
                
                # Get database size info
                self.cursor.execute("""
                    SELECT pg_size_pretty(pg_total_relation_size('documents')),
                        pg_size_pretty(pg_total_relation_size('embeddings'))
                """)
                doc_size, emb_size = self.cursor.fetchone()
                
                stats.update({
                    "document_count": doc_count,
                    "embedding_count": emb_count,
                    "documents_table_size": doc_size,
                    "embeddings_table_size": emb_size
                })
            except Exception as e:
                stats["error"] = f"Error retrieving PostgreSQL stats: {str(e)}"
        
        return stats    
             
    def save(self, path: str = None):
        """
        Save the vector store to disk
        
        Args:
            path: Path to save the vector store
            
        Returns:
            Boolean indicating success
        """
        import pickle
        import os
        
        if path is None:
            path = f"{self.collection_name}_vector_store.pkl"
            
        try:
            if self.store_type == "in_memory":
                # Save in-memory data directly
                with open(path, "wb") as f:
                    pickle.dump({
                        "documents": self.documents,
                        "embeddings": self.embeddings, 
                        "metadatas": self.metadatas,
                        "ids": self.ids,
                        "dimension": self.dimension,
                        "collection_name": self.collection_name
                    }, f)
                return True
                
            elif self.store_type == "faiss":
                import faiss
                
                # Create directory if it doesn't exist
                os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
                
                # Save FAISS index
                index_path = f"{path}.index"
                faiss.write_index(self.index, index_path)
                
                # Save associated documents
                with open(f"{path}.docs", "wb") as f:
                    pickle.dump({
                        "documents": self.documents,
                        "vector_id": self.vector_id,
                        "dimension": self.dimension,
                        "collection_name": self.collection_name
                    }, f)
                return True
                
            elif self.store_type == "chroma":
                # ChromaDB has its own persistence
                if hasattr(self.collection, "persist"):
                    self.collection.persist()
                return True
                
            elif self.store_type == "postgres":
                # PostgreSQL already persists data
                return True
                
            else:
                return False
                
        except Exception as e:
            print(f"Error saving vector store: {str(e)}")
            return False
    
    @classmethod
    def load(cls, path: str, embedding_model=None):
        """
        Load a vector store from disk
        
        Args:
            path: Path to the saved vector store
            embedding_model: Embedding model to use
            
        Returns:
            VectorStore instance
        """
        import pickle
        import os
        
        try:
            # Check if this is a FAISS index
            if os.path.exists(f"{path}.index") and os.path.exists(f"{path}.docs"):
                import faiss
                
                # Load documents and metadata
                with open(f"{path}.docs", "rb") as f:
                    data = pickle.load(f)
                
                # Create instance
                vector_store = cls(
                    embedding_model=embedding_model,
                    store_type="faiss",
                    collection_name=data.get("collection_name", "documents"),
                    dimension=data.get("dimension", 1536)
                )
                
                # Load FAISS index
                vector_store.index = faiss.read_index(f"{path}.index")
                vector_store.documents = data["documents"]
                vector_store.vector_id = data.get("vector_id", len(data["documents"]))
                
                return vector_store
                
            # Otherwise assume it's an in-memory store
            else:
                with open(path, "rb") as f:
                    data = pickle.load(f)
                
                # Create instance
                vector_store = cls(
                    embedding_model=embedding_model,
                    store_type="in_memory",
                    collection_name=data.get("collection_name", "documents"),
                    dimension=data.get("dimension", 1536)
                )
                
                # Load data
                vector_store.documents = data["documents"]
                vector_store.embeddings = data["embeddings"]
                vector_store.metadatas = data["metadatas"]
                vector_store.ids = data["ids"]
                
                return vector_store
                
        except Exception as e:
            raise ValueError(f"Error loading vector store: {str(e)}")
    
    def add_texts(self, texts: List[str], metadatas: Optional[List[Dict[str, Any]]] = None) -> List[str]:
        """
        Add texts to the vector store with optional metadata
        
        Args:
            texts: List of text strings to add
            metadatas: Optional list of metadata dictionaries
            
        Returns:
            List of document IDs
        """
        if metadatas is None:
            metadatas = [{} for _ in texts]
            
        # Create Document objects
        documents = []
        for i, text in enumerate(texts):
            metadata = metadatas[i] if i < len(metadatas) else {}
            doc = Document(page_content=text, metadata=metadata)
            documents.append(doc)
            
        # Add documents
        return self.add_documents(documents)
    
    def search(self, query: str, k: int = 4, metadata_filter: Optional[Dict[str, Any]] = None) -> List[Document]:
        """
        Search for documents similar to query with optional metadata filtering
        
        Args:
            query: Query string
            k: Number of documents to retrieve
            metadata_filter: Optional filter for document metadata
            
        Returns:
            List of Document objects
        """
        # First do similarity search
        results = self.similarity_search(query, k=k)
        
        # Apply metadata filter if specified
        if metadata_filter is not None:
            filtered_results = []
            for doc in results:
                match = True
                for key, value in metadata_filter.items():
                    if key not in doc.metadata or doc.metadata[key] != value:
                        match = False
                        break
                if match:
                    filtered_results.append(doc)
            return filtered_results
        else:
            return results
    
    def update_document(self, document_id: str, document: Document) -> bool:
        """
        Update a document in the vector store
        
        Args:
            document_id: ID of document to update
            document: New document content
            
        Returns:
            Boolean indicating success
        """
        # Delete the old document
        if not self.delete([document_id]):
            return False
            
        # Add the new document
        ids = self.add_documents([document])
        return len(ids) > 0
    
    def get_document(self, document_id: str) -> Optional[Document]:
        """
        Retrieve a document by ID
        
        Args:
            document_id: Document ID
            
        Returns:
            Document or None if not found
        """
        if self.store_type == "in_memory":
            for i, doc_id in enumerate(self.ids):
                if doc_id == document_id:
                    return self.documents[i]
            return None
            
        elif self.store_type == "chroma":
            try:
                result = self.collection.get(ids=[document_id])
                if result and result["documents"]:
                    return Document(
                        page_content=result["documents"][0],
                        metadata=result["metadatas"][0] if result["metadatas"] else {}
                    )
                return None
            except:
                return None
                
        elif self.store_type == "postgres":
            try:
                self.cursor.execute("""
                    SELECT content, metadata FROM documents WHERE id = %s
                """, (int(document_id),))
                
                result = self.cursor.fetchone()
                if result:
                    return Document(
                        page_content=result[0],
                        metadata=json.loads(result[1]) if result[1] else {}
                    )
                return None
            except:
                return None
                
        elif self.store_type == "faiss":
            # FAISS doesn't support direct document lookup by ID
            # We would need to iterate through the documents
            for doc in self.documents:
                if doc.metadata.get('faiss_id') == document_id:
                    return doc
            return None
    
    def __len__(self) -> int:
        """Return the number of documents in the store"""
        if self.store_type == "in_memory":
            return len(self.documents)
            
        elif self.store_type == "faiss":
            return len(self.documents)
            
        elif self.store_type == "chroma":
            try:
                return len(self.collection.get()["ids"])
            except:
                return 0
                
        elif self.store_type == "postgres":
            try:
                self.cursor.execute("SELECT COUNT(*) FROM documents")
                return self.cursor.fetchone()[0]
            except:
                return 0
    
    def __del__(self):
        """Clean up resources when the object is destroyed"""
        if self.store_type == "postgres" and hasattr(self, 'conn'):
            try:
                self.conn.close()
            except:
                pass    

    def retrieve(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Retrieve documents relevant to the query with scores and filtering options
        
        Args:
            query: Query string
            k: Number of documents to retrieve
            threshold: Optional similarity threshold (0.0-1.0)
            metadata_filter: Optional filter for document metadata
            
        Returns:
            List of dictionaries with document content, metadata, and similarity score
        """
        # Get query embedding
        query_embedding = self._get_embeddings([query])[0]
        
        results = []
        
        # Retrieval based on store type
        if self.store_type == "in_memory":
            import numpy as np
            
            if not self.embeddings:
                return []
                
            # Convert to numpy array for efficient computation
            embeddings_array = np.array(self.embeddings)
            query_array = np.array(query_embedding)
            
            # Compute similarities
            similarities = np.dot(embeddings_array, query_array) / (
                np.linalg.norm(embeddings_array, axis=1) * np.linalg.norm(query_array)
            )
            
            # Apply threshold if provided
            indices = list(range(len(similarities)))
            if threshold is not None:
                filtered_indices = [(i, sim) for i, sim in zip(indices, similarities) if sim >= threshold]
                indices, similarities = zip(*filtered_indices) if filtered_indices else ([], [])
                
            # Sort by similarity
            sorted_pairs = sorted(zip(indices, similarities), key=lambda x: x[1], reverse=True)
            
            # Apply metadata filter if provided
            filtered_indices = []
            for idx, score in sorted_pairs:
                if metadata_filter is None:
                    filtered_indices.append((idx, score))
                else:
                    doc_metadata = self.metadatas[idx] if idx < len(self.metadatas) else {}
                    match = True
                    for key, value in metadata_filter.items():
                        if key not in doc_metadata or doc_metadata[key] != value:
                            match = False
                            break
                    if match:
                        filtered_indices.append((idx, score))
                
                if len(filtered_indices) >= k:
                    break
            
            # Build result objects
            for idx, score in filtered_indices:
                results.append({
                    "document": self.documents[idx],
                    "content": self.documents[idx].page_content,
                    "metadata": self.documents[idx].metadata,
                    "score": float(score),
                    "id": self.ids[idx] if idx < len(self.ids) else None
                })
                
        elif self.store_type == "faiss":
            import numpy as np
            
            if len(self.documents) == 0:
                return []
                
            # Convert query embedding to numpy array
            query_array = np.array([query_embedding]).astype('float32')
            
            # Search FAISS index
            distances, indices = self.index.search(query_array, k if metadata_filter is None else min(len(self.documents), k*5))
            
            # Convert distances to similarity scores (FAISS returns L2 distance)
            max_distance = float(np.max(distances)) if distances.size > 0 else 1.0
            scores = 1.0 - distances[0] / (max_distance if max_distance > 0 else 1.0)
            
            # Apply metadata filter if provided
            count = 0
            for i, idx in enumerate(indices[0]):
                if idx < 0 or idx >= len(self.documents):
                    continue
                    
                if metadata_filter is not None:
                    doc_metadata = self.documents[idx].metadata
                    match = True
                    for key, value in metadata_filter.items():
                        if key not in doc_metadata or doc_metadata[key] != value:
                            match = False
                            break
                    if not match:
                        continue
                
                # Apply threshold if provided
                score = float(scores[i])
                if threshold is not None and score < threshold:
                    continue
                    
                results.append({
                    "document": self.documents[idx],
                    "content": self.documents[idx].page_content,
                    "metadata": self.documents[idx].metadata,
                    "score": score,
                    "id": self.documents[idx].metadata.get('faiss_id')
                })
                
                count += 1
                if count >= k:
                    break
                
        elif self.store_type == "chroma":
            # Query ChromaDB with include parameters
            chroma_results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=k if metadata_filter is None else min(1000, k*5),
                include=["documents", "metadatas", "distances", "embeddings"],
                where=metadata_filter  # ChromaDB supports direct metadata filtering
            )
            
            if not chroma_results["ids"][0]:
                return []
                
            # Convert distances to similarity scores (ChromaDB returns L2 distance)
            distances = chroma_results['distances'][0]
            max_distance = max(distances) if distances else 1.0
            scores = [1.0 - dist / (max_distance if max_distance > 0 else 1.0) for dist in distances]
            
            # Apply threshold if provided
            filtered_indices = list(range(len(scores)))
            if threshold is not None:
                filtered_indices = [i for i, score in enumerate(scores) if score >= threshold]
                
            # Build result objects (limited to k)
            for i in filtered_indices[:k]:
                results.append({
                    "document": Document(
                        page_content=chroma_results['documents'][0][i],
                        metadata=chroma_results['metadatas'][0][i] if chroma_results['metadatas'][0] else {}
                    ),
                    "content": chroma_results['documents'][0][i],
                    "metadata": chroma_results['metadatas'][0][i] if chroma_results['metadatas'][0] else {},
                    "score": scores[i],
                    "id": chroma_results['ids'][0][i]
                })
                
        elif self.store_type == "postgres":
            try:
                import psycopg2.extras
                import numpy as np
                
                # Convert query vector to PostgreSQL format
                query_vector_str = str(query_embedding).replace('[', '{').replace(']', '}')
                
                # Build WHERE clause for metadata filtering
                metadata_where = ""
                params = [query_vector_str, k]
                
                if metadata_filter:
                    metadata_conditions = []
                    for key, value in metadata_filter.items():
                        metadata_conditions.append(f"metadata->>{psycopg2.extensions.adapt(key)::text} = %s")
                        params.append(str(value))
                        
                    if metadata_conditions:
                        metadata_where = "AND " + " AND ".join(metadata_conditions)
                
                try:
                    # Try using vector similarity if extension is available
                    self.cursor.execute(f"""
                        SELECT d.id, d.content, d.metadata, 1 - (e.vector <-> %s::float[]) / 2 as similarity 
                        FROM documents d
                        JOIN embeddings e ON d.embedding_id = e.id
                        WHERE 1=1 {metadata_where}
                        ORDER BY similarity DESC
                        LIMIT %s
                    """, tuple(params))
                except:
                    # Fallback if vector extension not available
                    print("Warning: PostgreSQL vector extension not available, using random order")
                    if metadata_filter:
                        self.cursor.execute(f"""
                            SELECT d.id, d.content, d.metadata
                            FROM documents d
                            WHERE 1=1 {metadata_where}
                            LIMIT %s
                        """, tuple(params[1:]))
                    else:
                        self.cursor.execute("""
                            SELECT d.id, d.content, d.metadata
                            FROM documents d
                            LIMIT %s
                        """, (k,))
                        
                rows = self.cursor.fetchall()
                
                # Process results
                for row in rows:
                    doc_id, content, metadata_json = row[:3]
                    score = float(row[3]) if len(row) > 3 else 0.5  # Default score if not available
                    
                    # Apply threshold if provided
                    if threshold is not None and score < threshold:
                        continue
                        
                    metadata = json.loads(metadata_json) if metadata_json else {}
                    
                    results.append({
                        "document": Document(page_content=content, metadata=metadata),
                        "content": content,
                        "metadata": metadata,
                        "score": score,
                        "id": str(doc_id)
                    })
                    
            except Exception as e:
                print(f"Error in PostgreSQL retrieval: {str(e)}")
                
        # Sort results by score
        results.sort(key=lambda x: x["score"], reverse=True)
        
        return results[:k]

    def retrieve_text(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[str]:
        """
        Retrieve document content relevant to the query with filtering options"""
        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return [result["content"] for result in results]
    
    def retrieve_metadata(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Retrieve metadata relevant to the query with filtering options"""
        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return [result["metadata"] for result in results]
    
    def retrieve_ids(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[str]:
        """
        Retrieve document IDs relevant to the query with filtering options"""
        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return [result["id"] for result in results]
    
    def retrieve_scores(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[float]:
        """
        Retrieve similarity scores relevant to the query with filtering options"""
        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return [result["score"] for result in results]
    
    def retrieve_documents(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[Document]:
        """
        Retrieve Document objects relevant to the query with filtering options"""
        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return [result["document"] for result in results]
    
    def retrieve_all(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Retrieve all information relevant to the query with filtering options"""
        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return results
    
    def retrieve_all_texts(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[str]:
        """
        Retrieve all document content relevant to the query with filtering options"""
        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return [result["content"] for result in results]
    
    def retrieve_all_metadata(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Retrieve all metadata relevant to the query with filtering options"""
        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return [result["metadata"] for result in results]
    
    def retrieve_all_ids(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[str]:
        """
        Retrieve all document IDs relevant to the query with filtering options"""
        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return [result["id"] for result in results]
    
    def retrieve_all_scores(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[float]:
        """
        Retrieve all similarity scores relevant to the query with filtering options"""
        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return [result["score"] for result in results]
    
    def retrieve_all_documents(self, query: str, k: int = 4, threshold: float = None, metadata_filter: Dict[str, Any] = None) -> List[Document]:

        results = self.retrieve(query, k=k, threshold=threshold, metadata_filter=metadata_filter)
        return [result["document"] for result in results]
    
# Test the VectorStore class
def test_vector_store():
    # from langchain import SentenceTransformerWrapper
    import os
    
    # Load a pre-trained Sentence Transformer model
    # model = SentenceTransformerWrapper("msmarco-distilbert-base-v2")
    
    # Initialize a vector store
    store = VectorStore(
        embedding_model=embeddings_model,
        store_type="faiss",
        collection_name="documents",
        dimension=embeddings_model.get_dimension()
    )
    
    # Add some documents
    store.add_texts([
        "This is a test document",
        "Another test document",
        "Yet another test document"
    ])
    store.add_documents([
        Document(page_content="Test document 1", metadata={"type": "test"}),
        Document(page_content="Test document 4", metadata={"type": "test1"}),
        Document(page_content="Test document 5", metadata={"type": "test2"}),
        Document(page_content="Test document 6", metadata={"type": "test3"})
    ])
    
    # Search for similar documents
    results = store.search("test", k=2)
    assert len(results) == 2
    
    # Save and load the vector store
    store.save("test_store.pkl")
    loaded_store = VectorStore.load("test_store.pkl", embedding_model=embeddings_model)
    
    # Check if the loaded store has the same documents
    # assert len(loaded_store) == 3
    
    # Clear the store
    # loaded_store.clear()
    # assert len(loaded_store) == 0
    a=loaded_store.retrieve_metadata("test1", k=2)
    print(a)
    # Retrieve different types of data
    c=loaded_store.retrieve_all("test", k=2)
    print(c)
    d=loaded_store.retrieve_all_documents("test2", k=2)
    print(d)
    e=loaded_store.retrieve_all_ids("test3", k=2)
    print(e)
    f=loaded_store.retrieve_all_metadata("test", k=2)
    print(f)
    g=loaded_store.retrieve_all_scores("test1", k=2)
    print(g)
    # Delete the store file
    # os.remove("test_store.pkl")
    store = None
    
    print("VectorStore test passed")

# test_vector_store()



class DocumentLoader:
    """Loads documents from various sources into a standardized format"""
    
    def __init__(self, 
                 source_type: str = None, 
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200,
                 metadata_extraction: bool = True):
        """
        Initialize document loader
        
        Args:
            source_type: Type of source ('file', 'web', 'database', etc.)
            chunk_size: Default size for chunking documents
            chunk_overlap: Default overlap between chunks
            metadata_extraction: Whether to extract metadata from documents
        """
        self.source_type = source_type
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.metadata_extraction = metadata_extraction
        self.parsers = {}
        
        # Register default parsers
        self._register_default_parsers()
    
    def _register_default_parsers(self):
        """Register default document parsers based on available libraries"""
        # Text parser is always available
        self.parsers['txt'] = self._parse_text
        
        # Check for PDF support
        try:
            import PyPDF2
            self.parsers['pdf'] = self._parse_pdf
        except ImportError:
            pass
            
        # Check for DOCX support
        try:
            import docx
            self.parsers['docx'] = self._parse_docx
        except ImportError:
            pass
            
        # Check for HTML support
        try:
            import bs4
            self.parsers['html'] = self._parse_html
        except ImportError:
            pass
    
    def load(self, source, **kwargs):
        """
        Load documents from source
        
        Args:
            source: Source identifier (filepath, URL, etc.)
            **kwargs: Additional source-specific parameters
            
        Returns:
            List of Document objects
        """
        # Determine source type if not explicitly provided
        source_type = kwargs.get('source_type', self.source_type)
        if source_type is None:
            source_type = self._infer_source_type(source)
            
        # Handle different source types
        if source_type == 'file':
            return self.load_from_file(source, **kwargs)
        elif source_type == 'directory':
            return self.load_from_directory(source, **kwargs)
        elif source_type == 'web':
            return self.load_from_web(source, **kwargs)
        elif source_type == 'database':
            return self.load_from_database(source, **kwargs)
        else:
            raise ValueError(f"Unsupported source type: {source_type}")
    
    def load_from_file(self, filepath, **kwargs):
        """Load documents from a file"""
        import os
        
        # Get file extension
        _, ext = os.path.splitext(filepath)
        ext = ext.lower()[1:]  # Remove the dot and convert to lowercase
        
        # Check if we have a parser for this extension
        if ext in self.parsers:
            parser = self.parsers[ext]
            return parser(filepath, **kwargs)
        else:
            # Default to text parser
            return self._parse_text(filepath, **kwargs)
    
    def load_from_directory(self, directory_path, **kwargs):
        """Load documents from all files in a directory"""
        import os
        
        documents = []
        recursive = kwargs.get('recursive', False)
        
        if recursive:
            # Walk through all subdirectories
            for root, _, files in os.walk(directory_path):
                for filename in files:
                    filepath = os.path.join(root, filename)
                    try:
                        docs = self.load_from_file(filepath, **kwargs)
                        documents.extend(docs)
                    except Exception as e:
                        print(f"Error loading {filepath}: {str(e)}")
        else:
            # Only process files in the top directory
            for filename in os.listdir(directory_path):
                filepath = os.path.join(directory_path, filename)
                if os.path.isfile(filepath):
                    try:
                        docs = self.load_from_file(filepath, **kwargs)
                        documents.extend(docs)
                    except Exception as e:
                        print(f"Error loading {filepath}: {str(e)}")
        
        return documents
    
    def load_from_web(self, url, **kwargs):
        """Load documents from a web URL"""
        try:
            import requests
            from bs4 import BeautifulSoup
            
            # Get web content
            response = requests.get(url)
            response.raise_for_status()  # Raise exception for HTTP errors
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract title
            title = soup.title.string if soup.title else url
            
            # Extract main content (simplified approach)
            main_content = ''
            
            # Try to find main content container
            main_tags = ['article', 'main', 'div[role="main"]', '.main-content', '#content']
            content_element = None
            
            for tag in main_tags:
                if '[' in tag and '=' in tag:
                    # Handle attribute selector like div[role="main"]
                    tag_name, attr_selector = tag.split('[', 1)
                    attr_name, attr_value = attr_selector.rstrip(']').split('=')
                    attr_value = attr_value.strip('"\'')
                    elements = soup.find_all(tag_name, {attr_name: attr_value})
                elif tag.startswith('.'):
                    # Handle class selector
                    elements = soup.find_all(class_=tag[1:])
                elif tag.startswith('#'):
                    # Handle id selector
                    elements = soup.find_all(id=tag[1:])
                else:
                    # Handle tag selector
                    elements = soup.find_all(tag)
                
                if elements:
                    content_element = elements[0]
                    break
            
            # If no main content container found, use body
            if not content_element:
                content_element = soup.body
            
            if content_element:
                # Extract text
                main_content = content_element.get_text(separator='\n')
            
            # Create metadata
            metadata = {
                'source': url,
                'title': title,
                'date_retrieved': datetime.datetime.now().isoformat()
            }
            
            # Create document
            doc = Document(
                page_content=main_content,
                metadata=metadata
            )
            
            # Split into chunks if needed
            chunk_size = kwargs.get('chunk_size', self.chunk_size)
            chunk_overlap = kwargs.get('chunk_overlap', self.chunk_overlap)
            
            if chunk_size:
                return self._chunk_document(doc, chunk_size, chunk_overlap)
            else:
                return [doc]
                
        except ImportError:
            raise ImportError("Web loading requires 'requests' and 'bs4' packages")
        except Exception as e:
            raise ValueError(f"Error loading from web: {str(e)}")
    
    def load_from_database(self, connection_info, **kwargs):
        """Load documents from a database"""
        try:
            import sqlalchemy
            
            # Create engine
            if isinstance(connection_info, str):
                # Connection string provided
                engine = sqlalchemy.create_engine(connection_info)
            else:
                # Connection parameters provided
                engine = connection_info
                
            # Get query
            query = kwargs.get('query')
            if not query:
                raise ValueError("Database loading requires a 'query' parameter")
                
            # Execute query
            with engine.connect() as conn:
                result = conn.execute(sqlalchemy.text(query))
                
                documents = []
                for row in result:
                    # Determine content and metadata columns
                    content_col = kwargs.get('content_column', 'content')
                    id_col = kwargs.get('id_column')
                    
                    # Convert row to dict
                    if hasattr(row, '_asdict'):
                        row_dict = row._asdict()
                    else:
                        row_dict = dict(row)
                        
                    # Extract content
                    if content_col in row_dict:
                        content = row_dict[content_col]
                    else:
                        # If no specific content column, use first column
                        content = row_dict[list(row_dict.keys())[0]]
                        
                    # Extract metadata
                    metadata = {}
                    for key, value in row_dict.items():
                        if key != content_col:
                            metadata[key] = value
                            
                    # Add document
                    doc = Document(
                        page_content=content,
                        metadata=metadata
                    )
                    documents.append(doc)
                    
                return documents
                
        except ImportError:
            raise ImportError("Database loading requires 'sqlalchemy' package")
        except Exception as e:
            raise ValueError(f"Error loading from database: {str(e)}")
    
    def _parse_text(self, filepath, **kwargs):
        """Parse a text file"""
        try:
            encoding = kwargs.get('encoding', 'utf-8')
            
            with open(filepath, 'r', encoding=encoding) as f:
                text = f.read()
                
            # Create metadata
            metadata = {
                'source': filepath,
                'filetype': 'text'
            }
            
            # Create document
            doc = Document(
                page_content=text,
                metadata=metadata
            )
            
            # Split into chunks if needed
            chunk_size = kwargs.get('chunk_size', self.chunk_size)
            chunk_overlap = kwargs.get('chunk_overlap', self.chunk_overlap)
            
            if chunk_size:
                return self._chunk_document(doc, chunk_size, chunk_overlap)
            else:
                return [doc]
                
        except Exception as e:
            raise ValueError(f"Error parsing text file: {str(e)}")
    
    def _parse_pdf(self, filepath, **kwargs):
        """Parse a PDF file"""
        import PyPDF2
        import os
        from datetime import datetime
        
        try:
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                pages = []
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():  # Only add non-empty pages
                        pages.append(text)
                
                # Combine page content
                content = '\n\n'.join(pages)
                
                # Get PDF metadata
                info = pdf_reader.metadata
                
                  # Modified date parsing removed, no time parsing needed
                creation_date = None
              
                # Create metadata dict
                metadata = {
                    'source': filepath,
                    'filetype': 'pdf',
                    'pages': len(pdf_reader.pages),
                    'title': info.title if info and info.title else os.path.basename(filepath),
                    'author': info.author if info and info.author else None,
                    'creation_date': creation_date,
                }
                
                # Extract file metrics
                file_size = os.path.getsize(filepath)
                metadata['file_size'] = file_size
                
                # Create document
                doc = Document(
                    page_content=content,
                    metadata=metadata
                )
                
                # Split into chunks if needed
                chunk_size = kwargs.get('chunk_size', self.chunk_size)
                chunk_overlap = kwargs.get('chunk_overlap', self.chunk_overlap)
                
                if chunk_size:
                    return self._chunk_document(doc, chunk_size, chunk_overlap)
                else:
                    return [doc]
                    
        except Exception as e:
            raise ValueError(f"Error parsing PDF file: {str(e)}")
    
    def _parse_docx(self, filepath, **kwargs):
        """Parse a DOCX file"""
        import docx
        import os
        
        try:
            doc = docx.Document(filepath)
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Only add non-empty paragraphs
                    full_text.append(text)
                    
            # Join paragraphs with newlines
            content = '\n'.join(full_text)
            
            # Create metadata
            metadata = {
                'source': filepath,
                'filetype': 'docx',
                'title': os.path.basename(filepath),
            }
            
            # Extract file metrics
            file_size = os.path.getsize(filepath)
            metadata['file_size'] = file_size
            
            # Create document
            doc = Document(
                page_content=content,
                metadata=metadata
            )
            
            # Split into chunks if needed
            chunk_size = kwargs.get('chunk_size', self.chunk_size)
            chunk_overlap = kwargs.get('chunk_overlap', self.chunk_overlap)
            
            if chunk_size:
                return self._chunk_document(doc, chunk_size, chunk_overlap)
            else:
                return [doc]
                
        except Exception as e:
            raise ValueError(f"Error parsing DOCX file: {str(e)}")
    
    def _parse_html(self, filepath, **kwargs):
        """Parse an HTML file"""
        from bs4 import BeautifulSoup
        import os
        
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
                
            # Parse HTML
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
                
            # Get text
            text = soup.get_text(separator='\n')
            
            # Break into lines and remove leading/trailing space
            lines = (line.strip() for line in text.splitlines())
            # Break multi-headlines into a line each
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            # Remove blank lines
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Create metadata
            metadata = {
                'source': filepath,
                'filetype': 'html',
                'title': soup.title.string if soup.title else os.path.basename(filepath),
            }
            
            # Extract file metrics
            file_size = os.path.getsize(filepath)
            metadata['file_size'] = file_size
            
            # Create document
            doc = Document(
                page_content=text,
                metadata=metadata
            )
            
            # Split into chunks if needed
            chunk_size = kwargs.get('chunk_size', self.chunk_size)
            chunk_overlap = kwargs.get('chunk_overlap', self.chunk_overlap)
            
            if chunk_size:
                return self._chunk_document(doc, chunk_size, chunk_overlap)
            else:
                return [doc]
                
        except Exception as e:
            raise ValueError(f"Error parsing HTML file: {str(e)}")
    
    def _chunk_document(self,document, chunk_size, chunk_overlap):
        """Split document into smaller chunks"""
        text = document.page_content

        # Safety limit for very large documents
        max_length = 1_000_000  # 1 million characters
        if len(text) > max_length:
            print(f"Warning: Document truncated from {len(text)} to {max_length} characters")
            text = text[:max_length]

        # If text is too short, no need to chunk
        if len(text) <= chunk_size:
            return [document]

        chunks = []
        i = 0
        while i < len(text):
            # Calculate chunk end position
            chunk_end = min(i + chunk_size, len(text))

            # If not at the end, try to adjust to a break character
            if chunk_end < len(text):
                search_start = max(i + chunk_size // 2, chunk_end - 100)
                for j in range(chunk_end - 1, search_start, -1):
                    if text[j] in ('\n', '.'):
                        chunk_end = j + 1  # Include the break character
                        break

            chunk_text = text[i:chunk_end]
            if not chunk_text:
                break

            chunk_metadata = document.metadata.copy()
            chunk_metadata['chunk_index'] = len(chunks)
            chunk_metadata['chunk_start'] = i
            chunk_metadata['chunk_end'] = chunk_end

            chunks.append(Document(page_content=chunk_text, metadata=chunk_metadata))

            if chunk_end == len(text):
                break

            # Update i and guarantee advancement
            new_i = chunk_end - chunk_overlap
            if new_i <= i:
                new_i = i + 1
            i = new_i

        return chunks
    
    def _infer_source_type(self, source):
        """Infer source type from the source parameter"""
        import os



    # Try importing document handling packages
    try:
        from fpdf import FPDF
        from docx import Document as DocxDocument
    except ImportError:
        print("Warning: FPDF and python-docx packages not installed. Some document types may not be supported.")

    def _infer_source_type(self, source):
        """Infer source type from the source parameter"""
        # Check if it's a URL
        if isinstance(source, str) and (source.startswith('http://') or source.startswith('https://')):
            return 'web'
            
        # Check if it's a file path
        if isinstance(source, str) and os.path.exists(source):
            if os.path.isdir(source):
                return 'directory'
            else:
                return 'file'
                
        # Default to file
        return 'file'


def register_parser(self, extension, parser_function):
        """
        Register a custom parser for a file extension
        
        Args:
            extension: File extension (without the dot)
            parser_function: Function to parse files of this type
        """
        self.parsers[extension.lower()] = parser_function
import datetime
# Test the DocumentLoader class
def test_document_loader():
    loader = DocumentLoader()
    
    # Load a text file
    docs = loader.load("test.txt")
    assert len(docs) == 1
    assert docs[0].page_content == "This is a test document"
    assert docs[0].metadata['source'] == "test.txt"

    # Load a PDF file
    docs = loader.load("test.pdf")
    assert len(docs) == 1
    assert "This is a test PDF document" in docs[0].page_content
    assert docs[0].metadata['source'] == "test.pdf"
    assert docs[0].metadata['filetype'] == "pdf"
    assert docs[0].metadata['pages'] == 1

    # Load a DOCX file
    docs = loader.load("test.docx")
    assert len(docs) == 1
    assert "This is a test DOCX document" in docs[0].page_content
    assert docs[0].metadata['source'] == "test.docx"
    assert docs[0].metadata['filetype'] == "docx"

    # Load an HTML file
    docs = loader.load("test.html")
    assert len(docs) == 1
    assert "This is a test HTML document" in docs[0].page_content
    assert docs[0].metadata['source'] == "test.html"
    assert docs[0].metadata['filetype'] == "html"
    assert docs[0].metadata['title'] == "Test HTML Document"

    # Load a directory
    docs = loader.load("test_docs", source_type="directory")
    assert len(docs) == 3

    # Load a web page
    docs = loader.load("https://en.wikipedia.org/wiki/Python_(programming_language)")
    assert len(docs) == 1
    assert "Python is an interpreted, high-level and general-purpose programming language." in docs[0].page_content
    assert docs[0].metadata['source'] == "https://en.wikipedia.org/wiki/Python_(programming_language)"
    assert docs[0].metadata['filetype'] == "html"
    assert docs[0].metadata['title'] == "Python (programming language) - Wikipedia"
    
    print("DocumentLoader test passed")

# test_document_loader()

# create the test file test.txt with the content This is a test document
# create the test PDF file test.pdf with the content This is a test PDF document
# create the test DOCX file test.docx with the content This is a test DOCX document
# create the test HTML file test.html with the content This is a test HTML document
# create a directory test_docs with some text files inside
# Run the test_document_loader function to test the DocumentLoader class.
# import fpdf as FPDF
from fpdf import FPDF
# from docx import Document
def create_test_files():
    """Create test files for DocumentLoader testing"""
    
    # Create test directory if it doesn't exist
    test_dir = "test_docs"
    os.makedirs(test_dir, exist_ok=True)
    
    # Create test.txt
    with open(os.path.join(test_dir, "test.txt"), "w", encoding="utf-8") as f:
        f.write("This is a test document")
    
    # Create test.pdf
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="This is a test PDF document", ln=1, align="L")
    pdf.output(os.path.join(test_dir, "test.pdf"))
    
    # Create test.docx
    doc = DocxDocument()
    doc.add_paragraph("This is a test DOCX document")
    doc.save(os.path.join(test_dir, "test.docx"))
    
    # Create test.html
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Test HTML Document</title>
    </head>
    <body>
        <h1>This is a test HTML document</h1>
    </body>
    </html>
    """
    with open(os.path.join(test_dir, "test.html"), "w", encoding="utf-8") as f:
        f.write(html_content)
    
    # Create some additional text files in test_docs
    for i in range(3):
        with open(os.path.join(test_dir, f"sample_{i}.txt"), "w", encoding="utf-8") as f:
            f.write(f"This is sample document {i}")
    
    print(f"Test files created in {test_dir} directory")
def test_document_loader(txt_path,pdf_path,docx_path,directory_path,html_path):
    # Use a smaller chunk size for testing and limit overlap to improve performance
    loader = DocumentLoader(chunk_size=2, chunk_overlap=1)
    print('loader instance created')
    docs = loader.load(txt_path)
    print("loaded docs")
#     create_test_files()
# except Exception as e:
#     print(f"Error creating test files: {str(e)}")


# test_document_loader()

def test_document_loader1(txt_path,pdf_path,docx_path,directory_path,html_path):
    loader = DocumentLoader()
    print('loader instance created')
    # Load a text file
    docs = loader.load(txt_path)
    print("loaded docs")
    assert len(docs) == 1
    assert docs[0].page_content == "This is a test document"
    assert docs[0].metadata['source'] == txt_path
    print(docs[0].metadata)

    # Load a PDF file
    docs = loader.load(pdf_path)
    assert len(docs) == 1
    assert "This is a test PDF document" in docs[0].page_content
    assert docs[0].metadata['source'] == pdf_path
    assert docs[0].metadata['filetype'] == "pdf"
    assert docs[0].metadata['pages'] == 1
    print(docs[0].metadata)

    # Load a DOCX file
    docs = loader.load(docx_path)
    assert len(docs) == 1
    assert "This is a test DOCX document" in docs[0].page_content
    assert docs[0].metadata['source'] == docx_path
    assert docs[0].metadata['filetype'] == "docx"
    print(docs[0].metadata)

    # Load an HTML file
    docs = loader.load(html_path)
    assert len(docs) == 1
    assert "This is a test HTML document" in docs[0].page_content
    assert docs[0].metadata['source'] == html_path
    assert docs[0].metadata['filetype'] == "html"
    assert docs[0].metadata['title'] == "Test HTML Document"
    print(docs[0].metadata)

    # Load a directory
    docs = loader.load(directory_path, source_type="directory")
    print(len(docs))
    # assert len(docs) == 7

    loader = DocumentLoader(chunk_size=1000000, chunk_overlap=0)
    # Load a web page
    docs = loader.load("https://en.wikipedia.org/wiki/Python_(programming_language)")
    print(len(docs
          ))
    assert len(docs) == 1
    print(docs[0].page_content)
    # assert "Python is an interpreted, high-level and general-purpose programming language." in docs[0].page_content
    assert docs[0].metadata['source'] == "https://en.wikipedia.org/wiki/Python_(programming_language)"
    print(docs[0].metadata)
    assert docs[0].metadata['title'] == "Python (programming language) - Wikipedia"
    
    print("DocumentLoader test passed")
dir=r"test_docs"
txt_path=r"test_docs\test.txt"
pdf_path=r"test_docs\test.pdf"
docx_path=r"test_docs\test.docx"
html_path=r"test_docs\test.html"
test_document_loader1(txt_path,pdf_path,docx_path,dir,html_path)

#a=chunk_document(document=Document(page_content="This is a test document", metadata={"source": "test.txt"}), chunk_size=10, chunk_overlap=2)
# print(a)
# Test the chunk_document function
def test_chunk_document():
    # Test chunking a simple document
    doc = Document(page_content="This is a test document", metadata={"source": "test.txt"})
    chunks = chunk_document(doc, chunk_size=4, chunk_overlap=0)
    print('chunks:',chunks)
    assert len(chunks) == 4
    assert chunks[0].page_content == "This "
    assert chunks[1].page_content == "is a "
    assert chunks[2].page_content == "test "
    assert chunks[3].page_content == "document"

    # Test chunking a longer document
    doc = Document(page_content="This is a longer test document with more content", metadata={"source": "test.txt"})
    chunks = chunk_document(doc, chunk_size=10, chunk_overlap=3)
    assert len(chunks) == 6
    assert chunks[0].page_content == "This is a "
    assert chunks[1].page_content == "longer te"
    assert chunks[2].page_content == "st docume"
    assert chunks[3].page_content == "nt with m"
    assert chunks[4].page_content == "ore conte"
    assert chunks[5].page_content == "nt"

    print("chunk_document test passed")


# test_chunk_document()


